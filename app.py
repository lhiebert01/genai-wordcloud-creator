import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from wordcloud import WordCloud, STOPWORDS
import numpy as np
from PIL import Image
import io
import base64
import re
import string
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from collections import Counter
import PyPDF2
import docx
import time
import os
from dotenv import load_dotenv
import openai

# Load environment variables
load_dotenv()

# Set page configuration to use wide mode and hide the Streamlit menu
st.set_page_config(
    page_title="GenAI Word Cloud Creator",
    page_icon="☁️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Hide Streamlit branding
hide_streamlit_style = """
    <style>
        #MainMenu {visibility: hidden;}
        footer {visibility: hidden;}
        .stDeployButton {display:none;}
        .css-18e3th9 {padding-top: 0rem; padding-bottom: 0rem;}
        .css-1d391kg {padding-top: 1rem; padding-bottom: 1rem;}
        .block-container {padding-top: 1rem; padding-bottom: 1rem;}
    </style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

# Download NLTK resources if not already downloaded
try:
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('punkt')
    nltk.download('stopwords')

# Get OpenAI API key - first try from secrets, then from environment variables
try:
    openai_api_key = st.secrets["openai"]["api_key"]
except:
    openai_api_key = os.getenv("OPENAI_API_KEY")

if not openai_api_key:
    st.error("OpenAI API key not found. Please check your .env file or Streamlit secrets.")

# Initialize session state variables
if 'processed_document_text' not in st.session_state:
    st.session_state.processed_document_text = None

# Custom CSS for better styling
st.markdown("""
<style>
    /* Improve button visibility */
    .stButton > button {
        background-color: #4CAF50;
        color: white;
        font-weight: bold;
        border-radius: 8px;
        padding: 0.5rem 1rem;
        border: none;
        box-shadow: 0 2px 5px rgba(0,0,0,0.2);
    }
    .stButton > button:hover {
        background-color: #45a049;
        box-shadow: 0 4px 8px rgba(0,0,0,0.3);
    }
    
    /* Improve container visibility */
    .css-1r6slb0, .css-12oz5g7 {
        border: 1px solid #e0e0e0;
        border-radius: 10px;
        padding: 10px;
        box-shadow: 0 2px 5px rgba(0,0,0,0.1);
    }
    
    /* Improve text area visibility */
    .stTextArea textarea {
        border: 2px solid #4CAF50;
        border-radius: 8px;
    }
    
    /* Improve response area for both light and dark modes */
    .response-area {
        background-color: #f0f2f6;
        color: #262730;
        padding: 20px;
        border-radius: 10px;
        border: 1px solid #e0e0e0;
        max-height: 400px;
        overflow-y: auto;
    }
    
    /* Dark mode adjustments */
    @media (prefers-color-scheme: dark) {
        .response-area {
            background-color: #262730;
            color: #ffffff;
            border: 1px solid #4a4a4a;
        }
    }
    
    /* Title styling */
    .title-container {
        display: flex;
        align-items: center;
        margin-bottom: 0;
    }
    
    .title-icon {
        font-size: 2.5rem;
        margin-right: 10px;
    }
    
    .attribution {
        font-size: 0.9rem;
        margin-top: -10px;
        margin-bottom: 20px;
        color: #666;
    }
    
    .attribution a {
        color: #4CAF50;
        text-decoration: none;
    }
    
    .attribution a:hover {
        text-decoration: underline;
    }
</style>
""", unsafe_allow_html=True)

# Title with icon and attribution
st.markdown('<div class="title-container"><span class="title-icon">☁️</span><h1>GenAI Word Cloud Creator</h1></div>', unsafe_allow_html=True)
st.markdown('<p class="attribution">Developed by <a href="https://www.linkedin.com/in/lindsayhiebert/" target="_blank">Lindsay Hiebert</a></p>', unsafe_allow_html=True)

st.write("Upload a document or use ChatGPT to generate text for a word cloud.")

# Initialize session state variables if they don't exist
if 'document_text' not in st.session_state:
    st.session_state.document_text = ""
if 'chatgpt_response' not in st.session_state:
    st.session_state.chatgpt_response = ""
if 'processed_chatgpt_text' not in st.session_state:
    st.session_state.processed_chatgpt_text = ""
if 'wordcloud_source' not in st.session_state:
    st.session_state.wordcloud_source = None
if 'uploaded_file_name' not in st.session_state:
    st.session_state.uploaded_file_name = ""
if 'last_action' not in st.session_state:
    st.session_state.last_action = None
if 'wc_width' not in st.session_state:
    st.session_state.wc_width = 800
if 'wc_height' not in st.session_state:
    st.session_state.wc_height = 800
if 'last_settings' not in st.session_state:
    st.session_state.last_settings = {}
if 'wordcloud_image' not in st.session_state:
    st.session_state.wordcloud_image = None
if 'current_wordcloud' not in st.session_state:
    st.session_state.current_wordcloud = None
if 'current_wordcloud_text' not in st.session_state:
    st.session_state.current_wordcloud_text = ""
if 'current_source_text' not in st.session_state:
    st.session_state.current_source_text = ""
if 'word_frequencies' not in st.session_state:
    st.session_state.word_frequencies = []
if 'current_tab' not in st.session_state:
    st.session_state.current_tab = "document"

# Function to check if settings have changed
def settings_changed():
    current_settings = {
        'max_words': max_words,
        'color_map': color_map,
        'background_color': background_color,
        'width': st.session_state.wc_width,
        'height': st.session_state.wc_height
    }
    
    if 'last_settings' not in st.session_state:
        st.session_state.last_settings = current_settings
        return True
    
    changed = current_settings != st.session_state.last_settings
    st.session_state.last_settings = current_settings
    return changed

def extract_text_from_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_docx(file):
    text = docx2txt.process(file)
    return text

def extract_text_from_txt(file):
    return file.getvalue().decode("utf-8")

def preprocess_text(text):
    # Remove special characters, numbers, and extra whitespace
    text = re.sub(r'[^\w\s]', '', text)
    text = re.sub(r'\d+', '', text)
    text = re.sub(r'\s+', ' ', text).strip()
    
    # Convert to lowercase
    text = text.lower()
    
    # Remove stopwords
    stop_words = set(stopwords.words('english'))
    words = text.split()
    filtered_words = [word for word in words if word not in stop_words and len(word) > 2]
    
    return ' '.join(filtered_words)

def generate_word_cloud(text, max_words=100, width=800, height=400, colormap='viridis', 
                        background_color='white'):
    
    # Create word cloud
    wordcloud = WordCloud(
        width=width,
        height=height,
        max_words=max_words,
        background_color=background_color,
        colormap=colormap,
        prefer_horizontal=0.9,
        collocations=True,
        min_font_size=4,
        mode="RGB"
    ).generate(text)
    
    # Save the wordcloud image to session state
    st.session_state.wordcloud_image = wordcloud.to_array()
    st.session_state.current_wordcloud = wordcloud
    
    return wordcloud

def process_text_once(text):
    """Process text once and store the results for reuse"""
    if not text:
        return None
    
    # Check if we've already processed this text
    if st.session_state.current_wordcloud_text == text:
        return text  # Return the text, not the wordcloud object
    
    # Process new text
    processed_text = preprocess_text(text)
    
    # Store word frequencies for reuse
    st.session_state.word_frequencies = get_all_words(processed_text, max_words=400)
    st.session_state.current_wordcloud_text = text
    
    return processed_text

def get_all_words(text, max_words=400):
    words = text.split()
    word_counts = Counter(words)
    return word_counts.most_common(max_words)

def display_word_cloud(text, max_words=100, width=800, height=400, colormap='viridis', 
                      background_color='white', source_text="Document"):
    
    if not text:
        st.warning("Please enter some text or upload a document to generate a word cloud.")
        return
    
    try:
        # Process text only once
        processed_text = process_text_once(text)
        
        # Generate word cloud with current settings
        wordcloud = generate_word_cloud(
            processed_text, max_words, width, height, colormap, background_color
        )
        
        # Store current source text
        st.session_state.current_source_text = source_text
        
        # Create columns for display
        main_col1, main_col2 = st.columns([3, 2])
        
        # Display word cloud
        with main_col1:
            st.subheader("Word Cloud")
            st.caption(f"Generated from: {source_text}")
            
            # Create a figure for the word cloud
            fig, ax = plt.subplots(figsize=(width/100, height/100))
            
            # Display the word cloud image
            ax.imshow(wordcloud.to_array(), interpolation='bilinear')
            ax.axis('off')
            
            # Display the figure
            st.pyplot(fig)
            
            # Save the figure to a buffer for download
            download_buf = BytesIO()
            fig.savefig(download_buf, format='png', dpi=300, bbox_inches='tight')
            download_buf.seek(0)
            
            # Use a container to prevent the download button from affecting the display
            download_container = st.container()
            with download_container:
                # Use a unique key based on timestamp to prevent conflicts
                unique_key = f"download_wordcloud_{hash(source_text)}_{int(time.time())}"
                st.download_button(
                    label="Download Word Cloud as PNG",
                    data=download_buf,
                    file_name=f"wordcloud_{source_text.replace(' ', '_').lower()}.png",
                    mime="image/png",
                    key=unique_key
                )
        
        # Display word frequency
        with main_col2:
            st.subheader("Word Frequency")
            
            # Use stored word frequencies
            all_words = st.session_state.word_frequencies
            
            # Create DataFrame for display
            df = pd.DataFrame(all_words, columns=["Word", "Count"])
            
            # Display with pagination
            st.dataframe(df, use_container_width=True, height=400)
            
            # Create separate containers for download buttons
            csv_download = st.container()
            txt_download = st.container()
            
            # Download word frequency data
            csv = df.to_csv(index=False)
            with csv_download:
                # Use unique keys for each download button
                csv_key = f"download_csv_{hash(source_text)}_{int(time.time())}"
                st.download_button(
                    label="Download Word Frequency CSV",
                    data=csv,
                    file_name=f"word_frequency_{source_text.replace(' ', '_').lower()}.csv",
                    mime="text/csv",
                    key=csv_key
                )
            
            # Download as TXT file with all words
            txt_content = "\n".join([f"{word}: {count}" for word, count in all_words])
            with txt_download:
                # Use unique keys for each download button
                txt_key = f"download_txt_{hash(source_text)}_{int(time.time())}"
                st.download_button(
                    label="Download Word Frequency TXT",
                    data=txt_content,
                    file_name=f"word_frequency_{source_text.replace(' ', '_').lower()}.txt",
                    mime="text/plain",
                    key=txt_key
                )
    except Exception as e:
        st.error(f"Error generating word cloud: {str(e)}")
        import traceback
        st.error(traceback.format_exc())

def save_to_docx(text, filename):
    doc = Document()
    doc.add_heading('ChatGPT Response', 0)
    
    paragraphs = text.split('\n\n')
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para)
    
    # Save the document
    doc_path = f"{filename}.docx"
    doc.save(doc_path)
    
    # Read the file to create a download button
    with open(doc_path, "rb") as file:
        return file.read()

def get_chatgpt_response(prompt):
    """Get response from ChatGPT API."""
    try:
        import openai
        openai.api_key = openai_api_key
        
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1000,
            temperature=0.7
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"Error getting response from ChatGPT: {str(e)}")
        return None

# Create sidebar for customization
with st.sidebar:
    st.title("Word Cloud Settings")
    
    # Create two columns for a more compact layout
    col1, col2 = st.columns(2)
    
    # Column 1 controls
    with col1:
        # Color map selection
        color_map = st.selectbox(
            "Color Scheme",
            ["viridis", "plasma", "inferno", "magma", "cividis", 
             "Greys", "Blues", "Reds", "Greens", "Purples", "Oranges",
             "rainbow", "jet", "turbo", "cool", "hot"],
            key="colormap"
        )
    
    # Column 2 controls
    with col2:
        # Max words
        max_words = st.slider("Max Words", 50, 500, 200, 10, key="max_words")
        
        # Background color
        background_color = st.color_picker("Background", "#FFFFFF", key="bg_color")
    
    # Resolution options
    st.subheader("Resolution Settings")
    
    # Aspect ratio selection
    aspect_ratio = st.selectbox(
        "Aspect Ratio",
        ["16:9", "9:16", "8:10", "10:8", "5:7", "7:5", "6:4", "4:6", "10:4", "10:3", "Square (1:1)", "Custom"],
        key="aspect_ratio"
    )
    
    # Resolution presets
    resolution_preset = st.selectbox(
        "Resolution Preset",
        ["HD (1280x720)", "Full HD (1920x1080)", "2K (2560x1440)", "4K (3840x2160)", "Custom"],
        key="resolution_preset"
    )
    
    # Set width and height based on selections
    if aspect_ratio == "Custom" or resolution_preset == "Custom":
        # Custom dimensions
        width_col, height_col = st.columns(2)
        with width_col:
            width = st.number_input("Width (px)", 400, 4000, 800, 100, key="custom_width")
        with height_col:
            height = st.number_input("Height (px)", 400, 4000, 800, 100, key="custom_height")
    else:
        # Calculate dimensions based on preset and aspect ratio
        if resolution_preset == "HD (1280x720)":
            base_width, base_height = 1280, 720
        elif resolution_preset == "Full HD (1920x1080)":
            base_width, base_height = 1920, 1080
        elif resolution_preset == "2K (2560x1440)":
            base_width, base_height = 2560, 1440
        elif resolution_preset == "4K (3840x2160)":
            base_width, base_height = 3840, 2160
        else:
            base_width, base_height = 1280, 720
        
        # Adjust for aspect ratio
        if aspect_ratio == "Square (1:1)":
            size = min(base_width, base_height)
            width, height = size, size
        elif aspect_ratio == "16:9":
            height = int(base_width * 9 / 16)
            width = base_width
        elif aspect_ratio == "9:16":
            width = int(base_height * 9 / 16)
            height = base_height
        elif aspect_ratio == "8:10":
            height = int(base_width * 10 / 8)
            width = base_width
        elif aspect_ratio == "10:8":
            width = int(base_height * 10 / 8)
            height = base_height
        elif aspect_ratio == "5:7":
            height = int(base_width * 7 / 5)
            width = base_width
        elif aspect_ratio == "7:5":
            width = int(base_height * 7 / 5)
            height = base_height
        elif aspect_ratio == "6:4":
            height = int(base_width * 4 / 6)
            width = base_width
        elif aspect_ratio == "4:6":
            width = int(base_height * 4 / 6)
            height = base_height
        elif aspect_ratio == "10:4":
            height = int(base_width * 4 / 10)
            width = base_width
        elif aspect_ratio == "10:3":
            height = int(base_width * 3 / 10)
            width = base_width
        
        # Display the calculated dimensions
        st.caption(f"Dimensions: {width}x{height} pixels")
    
    # Store dimensions in session state
    st.session_state.wc_width = width
    st.session_state.wc_height = height

# Create tabs for different input methods
document_tab, chatgpt_tab = st.tabs(["Document Upload", "ChatGPT"])

# Store the current tab in session state
if 'current_tab' not in st.session_state:
    st.session_state.current_tab = "document"

# Document Upload Tab
with document_tab:
    st.session_state.current_tab = "document"
    
    # File uploader
    uploaded_file = st.file_uploader("Upload a document (PDF, DOCX, or TXT)", type=["pdf", "docx", "txt"])
    
    if uploaded_file:
        try:
            # Extract text based on file type
            if uploaded_file.type == "application/pdf":
                text = extract_text_from_pdf(uploaded_file)
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                text = extract_text_from_docx(uploaded_file)
            else:  # Assume TXT
                text = uploaded_file.getvalue().decode("utf-8")
            
            # Store the document text and file name in session state
            st.session_state.document_text = text
            st.session_state.uploaded_file_name = uploaded_file.name
            st.session_state.wordcloud_source = 'file'
            st.session_state.last_action = "upload"
            
            # Process the text for word cloud
            processed_text = preprocess_text(text)
            st.session_state.processed_document_text = processed_text
            
            # Display document info
            st.subheader("Document Information")
            st.write(f"**File Name:** {uploaded_file.name}")
            st.write(f"**File Size:** {uploaded_file.size / 1024:.2f} KB")
            st.write(f"**Word Count:** {len(text.split())}")
            
            # Display a sample of the text (not in an expander)
            st.subheader("Text Preview")
            st.text_area("Document Content (First 1000 characters)", 
                        value=text[:1000] + ("..." if len(text) > 1000 else ""), 
                        height=150, 
                        disabled=True)
            
            # Add a button to generate word cloud from document
            if st.button("Generate Word Cloud from Document", key="doc_generate_btn"):
                # Display word cloud
                display_word_cloud(
                    text=text,
                    max_words=max_words,
                    width=st.session_state.wc_width,
                    height=st.session_state.wc_height,
                    colormap=color_map,
                    background_color=background_color,
                    source_text=uploaded_file.name
                )
        except Exception as e:
            st.error(f"Error processing file: {str(e)}")
            import traceback
            st.error(traceback.format_exc())

# ChatGPT Tab
with chatgpt_tab:
    st.session_state.current_tab = "chatgpt"
    
    # Input for ChatGPT prompt
    prompt = st.text_area("Enter your prompt for ChatGPT", height=150)
    
    # Button to submit to ChatGPT
    chat_col1, chat_col2 = st.columns([1, 3])
    
    with chat_col1:
        # Use a unique key for this button
        if st.button("Submit to ChatGPT", key="submit_to_chatgpt_tab"):
            if prompt:
                with st.spinner("Generating response from ChatGPT..."):
                    response = get_chatgpt_response(prompt)
                    if response:
                        st.session_state.chatgpt_response = response
                        st.session_state.wordcloud_source = 'chat'
                        st.session_state.last_action = "chatgpt"
                        
                        # Process the text for word cloud
                        processed_text = preprocess_text(response)
                        st.session_state.processed_chatgpt_text = processed_text
            else:
                st.warning("Please enter a prompt for ChatGPT.")
    
    # Display ChatGPT response if available
    if st.session_state.get('chatgpt_response') and st.session_state.get('last_action') == "chatgpt":
        st.subheader("ChatGPT Response")
        
        # Display the response in an expander
        with st.expander("Response Text", expanded=True):
            # Create a text area for editing the response
            edited_response = st.text_area(
                "Edit response if needed:",
                value=st.session_state.chatgpt_response,
                height=300,
                key="edited_response"
            )
            
            # Create columns for download options
            download_col1, download_col2 = st.columns(2)
            
            with download_col1:
                # Download as TXT
                st.download_button(
                    label="Download as TXT",
                    data=edited_response,
                    file_name="chatgpt_response.txt",
                    mime="text/plain",
                    key="download_txt_response"
                )
            
            with download_col2:
                # Download as DOCX
                docx_data = save_to_docx(edited_response, "chatgpt_response")
                st.download_button(
                    label="Download as DOCX",
                    data=docx_data,
                    file_name="chatgpt_response.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_docx_response"
                )
            
            # Button to generate word cloud from edited response
            if st.button("Generate Word Cloud from ChatGPT Response", key="chatgpt_generate_btn"):
                # Use the edited response for the word cloud
                display_word_cloud(
                    text=edited_response,
                    max_words=max_words,
                    width=st.session_state.wc_width,
                    height=st.session_state.wc_height,
                    colormap=color_map,
                    background_color=background_color,
                    source_text="ChatGPT Response"
                )

# Display word cloud based on source
if st.session_state.wordcloud_source == 'file' and st.session_state.processed_document_text:
    display_word_cloud(
        st.session_state.processed_document_text, 
        max_words=max_words,
        width=st.session_state.wc_width,
        height=st.session_state.wc_height,
        colormap=color_map,
        background_color=background_color,
        source_text=st.session_state.uploaded_file_name
    )
elif st.session_state.wordcloud_source == 'chat' and st.session_state.processed_chatgpt_text:
    display_word_cloud(
        st.session_state.processed_chatgpt_text, 
        max_words=max_words,
        width=st.session_state.wc_width,
        height=st.session_state.wc_height,
        colormap=color_map,
        background_color=background_color,
        source_text="ChatGPT Response"
    )
else:
    # Display instructions
    st.info("Upload a document or use ChatGPT to generate text for a word cloud.")

# App instructions
st.subheader("How to Use")
st.markdown("""
1. **Upload a document** (PDF, DOCX, or TXT) or **ask ChatGPT** to generate text
2. The app **extracts** and **processes** the text
3. Choose a **color scheme and style** for your word cloud
4. A **word cloud** is generated based on word frequency
5. You can **customize** the appearance and download the result
""")

# Regenerate word cloud if settings have changed
if settings_changed():
    if st.session_state.current_wordcloud_text:
        display_word_cloud(
            st.session_state.current_wordcloud_text, 
            max_words=max_words,
            width=st.session_state.wc_width,
            height=st.session_state.wc_height,
            colormap=color_map,
            background_color=background_color,
            source_text=st.session_state.current_source_text
        )
